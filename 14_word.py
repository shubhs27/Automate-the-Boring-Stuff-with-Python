import docx

d= docx.Document('demo.docx')
print(d.paragraphs)
print()

print(d.paragraphs[0].text)
print()

# a new run starts whenever there is a change in style (bold, italics etc.)
p= d.paragraphs[1]
print(p.text)
print(p.runs[0].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[2].bold)

p.runs[4].underline=True
p.runs[4].text= 'italic and underlined'
d.save('demo1.docx')

print(p.style)      # Check all styles on word by ctrl+shift+alt+s
p.style= 'Title'
print()

#########################################################################

d= docx.Document()
d.add_paragraph('This paragraph was written from python code.')
d.add_paragraph('This too.')

p= d.paragraphs[1]
p.add_run('This is a new run.')
p.runs[1].bold=True
#d.save('demo2.docx')

#########################################################################

def getText(filename):
    doc= docx.Document(filename)
    fullText= []

    for para in doc.paragraphs:
        fullText.append(para.text)

    return '\n'.join(fullText)

print(getText('demo.docx'))